from transformers import pipeline
import re
from docx import Document
import PyPDF2
from summarization import generate_summary


def process_text_file(path:str):
    if path.endswith(".txt"):
        file_path = 'temp/a.txt'  
        with open(file_path, 'r') as file:
            file_content = file.read()
        summary = generate_summary(file_content)
        return summary
    
    elif path.endswith(".pdf"):
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            text = ""
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
        summary = generate_summary(text)
        return summary
    
    elif path.endswith(".docx"):
        doc = Document(path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
        summary = generate_summary(text)
        return summary

file_path = 'temp/a.txt'
file_path2 = "temp/test.pdf"
file_path3 = "temp/abcd.docx"

# print("\n")
# # print(get_summary(file_path))
# print("\n")
# print(process_text_file(file_path2))
# print("\n")
# print(get_summary(file_path))







# # Open the file in read mode
# file_path = 'temp/a.txt'  # replace 'your_file.txt' with the actual path to your text file
# with open(file_path, 'r') as file:
#     # Read the entire content of the file
#     file_content = file.read()

# filtered = clean_text(file_content)

# # Print or use the content as needed
# print(len(file_content))
# print(summarizer(filtered, max_length=40, do_sample=False))
# # print(summarizer(file_content))




# def extract_text_from_pdf(pdf_path):
#     # Open the PDF file in binary mode
#     with open(pdf_path, 'rb') as file:
#         # Create a PDF reader object
#         pdf_reader = PyPDF2.PdfReader(file)

#         # Get the number of pages in the PDF
#         num_pages = len(pdf_reader.pages)

#         # Initialize an empty string to store the extracted text
#         text = ""

#         # Iterate through all pages and extract text
#         for page_num in range(num_pages):
#             # Get the page object
#             page = pdf_reader.pages[page_num]

#             # Extract text from the page
#             text += page.extract_text()
#     print(text)
#     return summarizer(text, max_length=70, min_length=20, do_sample=False)

# # Example usage:
# pdf_path = 'temp/abcdef.pdf'  # Replace 'example.pdf' with the path to your PDF file
# extracted_text = extract_text_from_pdf(pdf_path)

# print(extracted_text)




# def extract_text_from_docx(docx_path):
#     # Create a Document object from the DOCX file
#     doc = Document(docx_path)

#     # Initialize an empty string to store the extracted text
#     text = ""

#     # Iterate through paragraphs and extract text
#     for paragraph in doc.paragraphs:
#         text += paragraph.text + '\n'

#     filtered = clean_text(text)
#     return summarizer(filtered, max_length=70, min_length=20, do_sample=False)

# # Example usage:
# docx_path = 'temp/abcd.docx'  # Replace 'example.docx' with the path to your DOCX file
# extracted_text = extract_text_from_docx(docx_path)

# print(extracted_text)

